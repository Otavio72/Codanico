import google.generativeai as genai
from docx import Document
import customtkinter as ctk
from tkinter import filedialog, messagebox
import threading
import os
from datetime import datetime
from dotenv import load_dotenv

# Carrega variáveis de ambiente do arquivo .env
load_dotenv()

# Obtém a API key do arquivo .env
API_KEY = os.getenv("API_KEY")

# Configura a API do Gemini
genai.configure(api_key=API_KEY)

# Define aparência padrão do CustomTkinter
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

# Garante que a pasta de códigos revisados existe
os.makedirs("CodigosRevisados", exist_ok=True)

# Prompt base para gerar documentação técnica
documentar = "Crie uma documentação técnica do código anteriormente mostrado. "
"Não repita o código, apenas explique: qual o propósito geral do programa, "
"as funções principais e como elas interagem. "
"Seja objetivo, conciso e profissional, escrevendo no máximo 5 parágrafos. "
"Não use listas, bullets, símbolos de Markdown ou docstrings. "
"Escreva como se fosse um relatório técnico para Word."


# Função para aplicar tema (claro/escuro)
def aplicar_tema(escolha):
    ctk.set_appearance_mode(escolha)


# Função que lê o conteúdo de um arquivo Python selecionado
def ler_arquivo(arquivoPython):
    try:
        with open(arquivoPython, 'r', encoding='utf-8') as arquivo:
            conteudo = arquivo.read()

        # Junta o conteúdo do arquivo em uma linha só (sem quebras de linha)
        inline = " ".join(line.strip() for line in conteudo.splitlines() if line.strip())

        # Chama o Codanico para processar o código
        Codanico(inline, arquivoPython)
    except Exception as e:
        # Exibe erro no terminal da interface
        terminal.delete("1.0", "end")
        terminal.insert("end", f"Ocorreu um Erro: {e}")


# Função para salvar o conteúdo do editor em um arquivo .py
def salvar_arquivo():
    try:
        caminho = filedialog.asksaveasfilename(defaultextension=".py", filetypes=[("Python Files", "*.py")])
        if caminho:
            conteudo = texto_editor.get("1.0", "end").strip()
            with open(caminho, "w", encoding="utf-8") as f:
                f.write(conteudo)
            messagebox.showinfo("Codanico", f"Arquivo salvo em:\n{caminho}")

    except Exception as e:
        terminal.delete("1.0", "end")
        terminal.insert("end", f"Ocorreu um Erro: {e}")


# Função para abrir um arquivo e enviar para o Codanico
def abrir_arquivo():
    try:
        # Abre janela para selecionar o arquivo
        arquivoPython = filedialog.askopenfilename(filetypes=[("Arquivos Python", "*.py")])
        if arquivoPython:
            # Cria thread para rodar o processamento sem travar a interface
            thread_Codanico = threading.Thread(target=ler_arquivo, args=(arquivoPython,))
            thread_Codanico.start()

            # Exibe conteúdo do arquivo na aba Editor
            with open(arquivoPython, "r", encoding="utf-8") as f:
                conteudo = f.read()
                texto_editor.delete("1.0", "end")
                texto_editor.insert("1.0", conteudo)

            # Atualiza título da janela
            gui.title(f"Codanico - {arquivoPython}")

            # Feedback no terminal
            terminal.delete("1.0", "end")
            terminal.insert("end", "Arquivo Recebido\n")
            terminal.insert("end", "Enviando para o Codanico...\n")
    except Exception as e:
        terminal.delete("1.0", "end")
        terminal.insert("end", f"Ocorreu um Erro: {e}")


# Função principal do Codanico — envia o código para o modelo de IA e gera comentários e documentação
def Codanico(codigoInline, arquivoPython):
    try:
        documento = Document()  # Cria novo documento Word

        # Inicializa o modelo do Gemini
        model = genai.GenerativeModel('gemini-2.0-flash')

        # Cria o chat com o contexto do papel do assistente
        chat = model.start_chat(history=[
            {
                'role': 'model',
                'parts': [
                    '''
                    Você é Codanico, um assistente especializado em revisar, comentar, corrigir e otimizar códigos Python.  

                    Sua função principal é adicionar **comentários curtos, claros e objetivos diretamente no código**, explicando cada parte importante de forma didática e funcional.  

                    APENAS quando solicitado, você deve gerar documentação do código; **nunca faça isso por conta própria**.  

                    Não use Markdown, não use docstrings e **não insira ```python``` ou ```** em nenhum momento.  
                    O código deve sair pronto para copiar e colar diretamente no editor.  

                    Sempre que necessário:
                    - Corrija **erros de sintaxe e runtime**.  
                    - Detecte e corrija **erros lógicos**.  
                    - Remova redundâncias, melhore a legibilidade e **otimize a lógica** sem alterar o comportamento.  
                    - Identifique problemas de **performance**, como loops desnecessários ou I/O ineficiente, e sugira melhorias comentadas no código.  
                    - Avise sobre **más práticas de manutenção**, como funções gigantes, uso excessivo de globals ou parâmetros demais.  
                    - Se houver **concorrência, threads ou acesso a dados compartilhados**, adicione comentários alertando sobre possíveis condições de corrida ou inconsistências.  

                    Não escreva explicações fora do código, exceto quando estiver **explicitamente gerando documentação solicitada**.
                    '''
                ],
            },
        ])

        # Envia o código Python para o modelo analisar e comentar
        resposta_codigo = chat.send_message(codigoInline)

        # Remove marcas de formatação de Markdown, caso venham
        codigo_limpo = resposta_codigo.text.replace("```python", "").replace("```", "").strip()

        # Gera nomes de arquivo baseados na data e hora
        data_hora = datetime.now().strftime("%Y-%m-%d_%H-%M")
        nome_base, extensao = os.path.splitext(os.path.basename(arquivoPython))

        nome_python = f"CodigosRevisados/{nome_base}_{data_hora}{extensao}"
        nome_doc = f"Documentações/Documentação_{nome_base}_{data_hora}"

        # Salva o código revisado
        with open(nome_python, 'w', encoding='utf-8') as arquivo:
            arquivo.write(codigo_limpo)

            # Gera documentação técnica usando o mesmo chat
            resposta_doc = chat.send_message(documentar)
            documentacaoLimpa = resposta_doc.text.replace("```python", "").replace("```", "").strip()

            # Monta o documento Word
            documento.add_heading(f'{nome_base}{extensao}', level=0)
            documento.add_paragraph(documentacaoLimpa)
            documento.save(f'{nome_doc}.docx')

            # Feedback visual no terminal
            terminal.insert("end", f"Código comentado salvo em: {nome_python}\n")
            terminal.insert("end", f"Documentação salva em: {nome_doc}\n")

    except Exception as e:
        # Mostra erros na interface
        terminal.delete("1.0", "end")
        terminal.insert("end", f"Ocorreu um Erro: {e}")


# ---------------------- INTERFACE GRÁFICA ----------------------

# Janela principal
gui = ctk.CTk()
gui.title("Codanico")
gui.geometry("800x600")

# Cria container de abas
tab_container = ctk.CTkTabview(gui)
tab_container.pack(fill="both", expand=True, padx=10, pady=10)

# Aba principal: Codanico
aba1 = tab_container.add("Codanico")

terminal = ctk.CTkTextbox(aba1, height=400, width=700)
terminal.pack(pady=20)

btn_enviar_arquivo = ctk.CTkButton(aba1, text="Enviar arquivo", command=abrir_arquivo)
btn_enviar_arquivo.pack(pady=10)

# Aba de edição de código
aba2 = tab_container.add("Editor")

texto_editor = ctk.CTkTextbox(aba2, height=400, width=700)
texto_editor.pack(pady=20)

frame_botoes = ctk.CTkFrame(aba2)
frame_botoes.pack(pady=10)

btn_salvar = ctk.CTkButton(frame_botoes, text="Salvar", command=salvar_arquivo)
btn_salvar.grid(row=0, column=1, padx=5)

# Aba de configurações
aba3 = tab_container.add("Configurações")

label_tema = ctk.CTkLabel(aba3, text="Tema da Interface:")
label_tema.pack(pady=10)

frame_tema = ctk.CTkFrame(aba3)
frame_tema.pack(pady=5)

btn_claro = ctk.CTkButton(frame_tema, text="Claro", command=lambda: aplicar_tema("light"))
btn_claro.grid(row=0, column=0, padx=10)

btn_escuro = ctk.CTkButton(frame_tema, text="Escuro", command=lambda: aplicar_tema("dark"))
btn_escuro.grid(row=0, column=1, padx=10)

# Inicia o loop da interface
gui.mainloop()
